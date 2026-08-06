# -*- coding: utf-8 -*-
"""Genera el documento de analisis y especificacion de requerimientos de BotaBien."""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from gen_doc_data import (FECHA, RESPONSABLE, PROYECTO, EQUIPO, VERSIONES, DEFINICIONES,
                          JUSTIFICACION, DOC_RELACIONADA, FLUJO_PROCESO, CASOS_USO,
                          RF, RNF, OBSERVACIONES, APROBACIONES)

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   "docs", "F_Analisis_de_Requerimientos_V1,0_BotaBien.docx")

CUS_IDS = [c[0] for c in CASOS_USO]

# ---------- helpers ----------

def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def style_cell(cell, text, bold=False, size=10, align=None, bg=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bg:
        set_cell_bg(cell, bg)


def add_table(doc, rows, cols, widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    return t


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.size = Pt(7)


def build_footer(section):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.text = ""
    # tabuladores: izquierda / centro / derecha
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    usable = section.page_width - section.left_margin - section.right_margin
    for pos, val in ((int(usable / 2), 'center'), (int(usable), 'right')):
        tab = OxmlElement('w:tab'); tab.set(qn('w:val'), val); tab.set(qn('w:pos'), str(pos))
        tabs.append(tab)
    pPr.append(tabs)
    r = p.add_run("Documento análisis y especificación de requerimientos  ")
    r.font.size = Pt(7)
    add_page_field(p)
    r2 = p.add_run("\t08-IF-019\tV.1")
    r2.font.size = Pt(7)


def h(doc, text, level, size=None, space_before=10):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        run.font.name = 'Calibri'
        if size:
            run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    return p


def para(doc, text, size=10, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic; r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, label, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(label); r1.bold = True; r1.font.size = Pt(size); r1.font.name = 'Calibri'
    r2 = p.add_run(text); r2.font.size = Pt(size); r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    return p


# ---------- trazabilidad ----------
def build_matrix():
    rows = []
    for rid, nombre, desc, actor, reglas, interop, rel, cus in RF:
        rows.append((rid, set(cus)))
    for rid, nombre, desc, prio, cus in RNF:
        rows.append((rid, set(cus)))
    return rows

MATRIX = build_matrix()
MATRIX_MAP = dict(MATRIX)

# requerimientos asociados por CUS, derivados de la matriz (garantiza consistencia)
def reqs_for_cus(cus_id):
    rfs = [rid for rid, s in MATRIX if rid.startswith('RF-') and cus_id in s]
    rnfs = [rid for rid, s in MATRIX if rid.startswith('RNF-') and cus_id in s]
    return ", ".join(rfs) + (", " + ", ".join(rnfs) if rnfs else "")


# ---------- documento ----------
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

sec = doc.sections[0]
sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
build_footer(sec)

# --- Portada ---
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(90)
r = p.add_run("DOCUMENTO DE ANÁLISIS Y ESPECIFICACIÓN DE REQUERIMIENTOS:")
r.bold = True; r.font.size = Pt(20); r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(PROYECTO)
r.bold = True; r.font.size = Pt(34); r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Clasificación de residuos por visión artificial en el dispositivo")
r.italic = True; r.font.size = Pt(12); r.font.name = 'Calibri'

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(150)
r = p.add_run("Versión 1,0   ·   %s   ·   %s" % (FECHA, RESPONSABLE))
r.font.size = Pt(11); r.font.name = 'Calibri'

doc.add_page_break()

# --- Tabla de contenido ---
h(doc, "Tabla de contenido", 1)
toc = ["1. Equipo del Proyecto", "2. Control de Versiones", "3. Definiciones, Siglas y Abreviaturas",
       "     3.1 Justificación de la necesidad", "4. Documentación relacionada",
       "5. Diagrama flujo de actividades del proceso", "6. Casos de Uso"]
for i, (cid, nombre, _a, _d, _f, _al) in enumerate(CASOS_USO, start=1):
    toc.append("     6.1.%d %s: %s" % (i, cid, nombre))
toc += ["7. Requerimientos", "     7.1 Requerimientos funcionales",
        "          7.1.1 Lista de requerimientos funcionales",
        "          7.1.2 Especificación de requerimientos funcionales",
        "     7.2 Requerimientos No Funcionales",
        "8. Matriz de trazabilidad de Requerimientos vs Casos de uso",
        "9. Observaciones adicionales",
        "10. Control de revisión y aprobaciones del documento y sus anexos"]
for line in toc:
    para(doc, line, size=10, space_after=2)

doc.add_page_break()

# --- 1. Equipo ---
h(doc, "1. Equipo del Proyecto", 1)
t = add_table(doc, 1 + len(EQUIPO), 2, widths=[6.0, 10.0])
style_cell(t.rows[0].cells[0], "Rol", bold=True, bg="D9E2F3")
style_cell(t.rows[0].cells[1], "Persona asignada", bold=True, bg="D9E2F3")
for i, (rol, persona) in enumerate(EQUIPO, start=1):
    style_cell(t.rows[i].cells[0], rol)
    style_cell(t.rows[i].cells[1], persona)

# --- 2. Control de versiones ---
h(doc, "2. Control de Versiones", 1)
t = add_table(doc, 1 + len(VERSIONES), 4, widths=[2.8, 2.0, 6.7, 4.5])
for j, head in enumerate(["Fecha", "Versión", "Descripción", "Responsable de la versión"]):
    style_cell(t.rows[0].cells[j], head, bold=True, bg="D9E2F3")
for i, fila in enumerate(VERSIONES, start=1):
    for j, val in enumerate(fila):
        style_cell(t.rows[i].cells[j], val)

# --- 3. Definiciones ---
h(doc, "3. Definiciones, Siglas y Abreviaturas", 1)
for termino, definicion in DEFINICIONES:
    bullet(doc, termino + ": ", definicion)

h(doc, "3.1 Justificación de la necesidad", 2)
for bloque in JUSTIFICACION.split("\n\n"):
    para(doc, bloque, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# --- 4. Documentación relacionada ---
h(doc, "4. Documentación relacionada", 1)
t = add_table(doc, 1 + len(DOC_RELACIONADA), 2, widths=[9.5, 6.5])
style_cell(t.rows[0].cells[0], "Título de documento", bold=True, bg="D9E2F3")
style_cell(t.rows[0].cells[1], "Ubicación", bold=True, bg="D9E2F3")
for i, (titulo, ubic) in enumerate(DOC_RELACIONADA, start=1):
    style_cell(t.rows[i].cells[0], titulo)
    style_cell(t.rows[i].cells[1], ubic)

# --- 5. Diagrama de flujo ---
h(doc, "5. Diagrama flujo de actividades del proceso", 1)
para(doc, "El proceso principal del sistema, desde la apertura de la aplicación hasta la entrega de la "
          "recomendación de caneca, se describe en la siguiente secuencia de actividades. El diagrama de flujo "
          "correspondiente, junto con los diagramas de casos de uso, de clases, de secuencia y de estados, se "
          "encuentra en el archivo docs/arquitectura.md del repositorio, en notación Mermaid.",
     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
for i, paso in enumerate(FLUJO_PROCESO, start=1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(paso); r.font.size = Pt(10); r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

# --- 6. Casos de uso ---
h(doc, "6. Casos de Uso", 1)
para(doc, "A continuación se especifican los casos de uso del sistema. Cada caso describe una interacción "
          "completa entre un actor y el sistema, con su flujo principal, sus flujos alternativos y los "
          "requerimientos que lo soportan.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

for i, (cid, nombre, actor, desc, flujo, alternos) in enumerate(CASOS_USO, start=1):
    h(doc, "6.1.%d  %s: %s" % (i, cid, nombre), 4, size=12)
    bullet(doc, "Actor: ", actor)
    bullet(doc, "Descripción: ", desc)
    p = doc.add_paragraph(); r = p.add_run("Flujo principal:")
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    for paso in flujo:
        pp = doc.add_paragraph(style='List Number')
        rr = pp.add_run(paso); rr.font.size = Pt(10); rr.font.name = 'Calibri'
        pp.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph(); r = p.add_run("Flujos alternativos:")
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)
    for alt in alternos:
        pp = doc.add_paragraph(style='List Bullet')
        rr = pp.add_run(alt); rr.font.size = Pt(10); rr.font.name = 'Calibri'
        pp.paragraph_format.space_after = Pt(0)
    bullet(doc, "Requerimientos asociados: ", reqs_for_cus(cid))

# --- 7. Requerimientos ---
doc.add_page_break()
h(doc, "7. Requerimientos", 1)
para(doc, "Un requerimiento describe una necesidad de negocio en términos de capacidades y/o servicios "
          "funcionales que ofrece un sistema, así como de las restricciones de calidad bajo las cuales debe "
          "operar. Los requerimientos funcionales definen qué debe hacer el sistema; los requerimientos no "
          "funcionales definen con qué nivel de calidad debe hacerlo.", align=WD_ALIGN_PARAGRAPH.JUSTIFY)

h(doc, "7.1 Requerimientos funcionales", 2)
h(doc, "7.1.1 Lista de requerimientos funcionales", 4, size=12)
t = add_table(doc, 1 + len(RF), 2, widths=[3.5, 12.5])
style_cell(t.rows[0].cells[0], "Identificador", bold=True, bg="D9E2F3")
style_cell(t.rows[0].cells[1], "Nombre requerimiento", bold=True, bg="D9E2F3")
for i, item in enumerate(RF, start=1):
    style_cell(t.rows[i].cells[0], item[0])
    style_cell(t.rows[i].cells[1], item[1])

h(doc, "7.1.2 Especificación de requerimientos funcionales", 4, size=12)
for rid, nombre, desc, actor, reglas, interop, rel, cus in RF:
    campos = [
        ("Identificador", rid),
        ("Nombre del Requerimiento", nombre),
        ("Descripción", desc),
        ("Actor", actor),
        ("Reglas de negocio relacionadas", reglas),
        ("Interoperabilidad con otro sistema, módulo o componente", interop),
        ("Relaciones entre requerimientos", rel),
        ("Casos de uso relacionados / Historias de usuario relacionados", ", ".join(cus)),
        ("Responsable elaboración", RESPONSABLE),
        ("Fecha de elaboración", FECHA),
    ]
    t = add_table(doc, len(campos), 2, widths=[5.5, 10.5])
    for i, (campo, valor) in enumerate(campos):
        style_cell(t.rows[i].cells[0], campo, bold=True, size=9, bg="EDF2FA")
        style_cell(t.rows[i].cells[1], valor, size=9)
    para(doc, "", size=6, space_after=8)

h(doc, "7.2 Requerimientos No Funcionales", 2)
t = add_table(doc, 1 + len(RNF), 2, widths=[3.5, 12.5])
style_cell(t.rows[0].cells[0], "Identificador", bold=True, bg="D9E2F3")
style_cell(t.rows[0].cells[1], "Nombre requerimiento", bold=True, bg="D9E2F3")
for i, item in enumerate(RNF, start=1):
    style_cell(t.rows[i].cells[0], item[0])
    style_cell(t.rows[i].cells[1], item[1])

para(doc, "", size=6, space_after=8)

for rid, nombre, desc, prio, cus in RNF:
    marcas = {"Alta": "Alta __X__   Media____   Baja____",
              "Media": "Alta ____   Media__X__   Baja____",
              "Baja": "Alta ____   Media____   Baja__X__"}[prio]
    campos = [
        ("Identificación del requerimiento", rid),
        ("Nombre del Requerimiento", nombre),
        ("Descripción", desc),
        ("Prioridad", marcas),
        ("Responsable elaboración", RESPONSABLE),
        ("Fecha de elaboración", FECHA),
    ]
    t = add_table(doc, len(campos), 2, widths=[5.5, 10.5])
    for i, (campo, valor) in enumerate(campos):
        style_cell(t.rows[i].cells[0], campo, bold=True, size=9, bg="EDF2FA")
        style_cell(t.rows[i].cells[1], valor, size=9)
    para(doc, "", size=6, space_after=8)

# --- 8. Matriz de trazabilidad (seccion apaisada) ---
land = doc.add_section(WD_SECTION.NEW_PAGE)
land.orientation = WD_ORIENT.LANDSCAPE
land.page_width, land.page_height = land.page_height, land.page_width
land.top_margin = Cm(1.8); land.bottom_margin = Cm(1.8)
land.left_margin = Cm(1.8); land.right_margin = Cm(1.8)
build_footer(land)

h(doc, "8. Matriz de trazabilidad de Requerimientos vs Casos de uso", 1)
para(doc, "RF: Requerimiento Funcional", size=9, space_after=0)
para(doc, "RNF: Requerimiento No Funcional", size=9, space_after=0)
para(doc, "CUS: Caso de Uso", size=9, space_after=8)

ncols = 1 + len(CUS_IDS)
widths = [3.2] + [2.3] * len(CUS_IDS)
t = add_table(doc, 1 + len(MATRIX), ncols, widths=widths)
style_cell(t.rows[0].cells[0], "", bold=True, bg="D9E2F3")
for j, cid in enumerate(CUS_IDS, start=1):
    style_cell(t.rows[0].cells[j], cid, bold=True, size=8,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg="D9E2F3")
for i, (rid, cusset) in enumerate(MATRIX, start=1):
    style_cell(t.rows[i].cells[0], rid, bold=True, size=8)
    for j, cid in enumerate(CUS_IDS, start=1):
        style_cell(t.rows[i].cells[j], "X" if cid in cusset else "", size=8,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

# --- volver a vertical ---
port = doc.add_section(WD_SECTION.NEW_PAGE)
port.orientation = WD_ORIENT.PORTRAIT
port.page_width, port.page_height = port.page_height, port.page_width
port.top_margin = Cm(2); port.bottom_margin = Cm(2)
port.left_margin = Cm(2.5); port.right_margin = Cm(2.5)
build_footer(port)

# --- 9. Observaciones ---
h(doc, "9. Observaciones adicionales", 1)
for bloque in OBSERVACIONES.split("\n\n"):
    para(doc, bloque, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

# --- 10. Control de revision ---
h(doc, "10. Control de revisión y aprobaciones del documento y sus anexos", 1)
t = add_table(doc, 1 + len(APROBACIONES), 4, widths=[4.5, 4.5, 3.0, 4.0])
for j, head in enumerate(["Rol", "Nombre", "Fecha", "Firma / Evidencia"]):
    style_cell(t.rows[0].cells[j], head, bold=True, bg="D9E2F3")
for i, fila in enumerate(APROBACIONES, start=1):
    for j, val in enumerate(fila):
        style_cell(t.rows[i].cells[j], val)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)

# ---------- verificacion ----------
esperado = 1 + 1 + 1 + 1 + len(RF) + 1 + len(RNF) + 1 + 1
d2 = Document(OUT)
print("Archivo:", OUT)
print("Tablas esperadas:", esperado, "| encontradas:", len(d2.tables))
print("CUS:", len(CASOS_USO), "| RF:", len(RF), "| RNF:", len(RNF))
vacios = [rid for rid, s in MATRIX if not s]
print("Filas de matriz vacías:", vacios if vacios else "ninguna")
sin_rf = [c for c in CUS_IDS if not any(c in s for rid, s in MATRIX if rid.startswith('RF-'))]
print("CUS sin RF asociado:", sin_rf if sin_rf else "ninguno")
ids_rf = [x[0] for x in RF]
ids_rnf = [x[0] for x in RNF]
print("Numeración RF continua:", ids_rf == ["RF-%03d" % k for k in range(1, len(RF) + 1)])
print("Numeración RNF continua:", ids_rnf == ["RNF-%03d" % k for k in range(1, len(RNF) + 1)])
print("Numeración CUS continua:", CUS_IDS == ["CUS-%03d" % k for k in range(1, len(CASOS_USO) + 1)])
